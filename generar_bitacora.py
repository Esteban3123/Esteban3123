#!/usr/bin/env python3
"""Genera el formato GFPI-F-147 Bitácora de Seguimiento Etapa Productiva SENA."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="000000", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)


def set_run_font(run, size=9, bold=False, color=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def write_cell(
    cell,
    text,
    *,
    size=8,
    bold=False,
    align="left",
    shade=None,
    color=None,
    clear=True,
):
    if clear:
        cell.text = ""
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if shade:
        set_cell_shading(cell, shade)
    set_cell_borders(cell)
    set_cell_vertical_align(cell, "center")
    return p


def merge_cells(table, r1, c1, r2, c2):
    cell = table.cell(r1, c1)
    cell.merge(table.cell(r2, c2))
    return cell


def add_heading_para(doc, text, size=11, bold=True, align="center", space_after=6):
    p = doc.add_paragraph()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


GREEN = "1F4E3D"
LIGHT_GREEN = "D9EAD3"
LIGHT_GRAY = "F3F3F3"
YELLOW = "FFF2CC"
WHITE = "FFFFFF"


def build_header(doc):
    table = doc.add_table(rows=4, cols=6)
    table.autofit = True

    # Fila 0: Código / Versión / Proceso
    merge_cells(table, 0, 0, 0, 1)
    write_cell(table.cell(0, 0), "Código:\nGFPI-F-147", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    merge_cells(table, 0, 2, 0, 3)
    write_cell(table.cell(0, 2), "Versión: 05", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    merge_cells(table, 0, 4, 0, 5)
    write_cell(table.cell(0, 4), "PROCESO\nGESTIÓN DE FORMACIÓN PROFESIONAL INTEGRAL", size=7, bold=True, align="center", shade=LIGHT_GREEN)

    # Fila 1: Nombre del formato
    merge_cells(table, 1, 0, 1, 5)
    write_cell(
        table.cell(1, 0),
        "NOMBRE DEL FORMATO\nFORMATO BITÁCORA DE SEGUIMIENTO ETAPA PRODUCTIVA",
        size=9,
        bold=True,
        align="center",
        shade=GREEN,
        color=(255, 255, 255),
    )

    # Fila 2: Clasificación
    merge_cells(table, 2, 0, 2, 5)
    write_cell(table.cell(2, 0), "CLASIFICACIÓN DE LA INFORMACIÓN", size=8, bold=True, align="center", shade=LIGHT_GRAY)

    # Fila 3: Tipos
    write_cell(table.cell(3, 0), "Pública", size=8, align="center")
    write_cell(table.cell(3, 1), "X", size=9, bold=True, align="center", shade=YELLOW)
    write_cell(table.cell(3, 2), "Pública Clasificada", size=8, align="center")
    write_cell(table.cell(3, 3), "", size=8, align="center")
    write_cell(table.cell(3, 4), "Pública Reservada", size=8, align="center")
    write_cell(table.cell(3, 5), "", size=8, align="center")

    doc.add_paragraph()


def build_bitacora_periodo(doc, datos):
    add_heading_para(doc, "Identificación de la bitácora", size=10, align="left", space_after=4)
    table = doc.add_table(rows=2, cols=4)
    write_cell(table.cell(0, 0), "Bitácora N°", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(table.cell(0, 1), datos.get("bitacora_n", ""), size=10, bold=True, align="center", shade=YELLOW)
    write_cell(table.cell(0, 2), "Período a reportar", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(table.cell(0, 3), "", size=8)
    write_cell(table.cell(1, 0), "Desde", size=8, bold=True, align="center", shade=LIGHT_GRAY)
    write_cell(table.cell(1, 1), datos.get("periodo_desde", ""), size=9, align="center", shade=YELLOW)
    write_cell(table.cell(1, 2), "Hasta", size=8, bold=True, align="center", shade=LIGHT_GRAY)
    write_cell(table.cell(1, 3), datos.get("periodo_hasta", ""), size=9, align="center", shade=YELLOW)
    doc.add_paragraph()


def build_datos_aprendiz(doc, d):
    add_heading_para(doc, "Datos del aprendiz", size=10, align="left", space_after=4)

    t1 = doc.add_table(rows=4, cols=4)
    labels_r0 = [
        ("Nombre completo del aprendiz", d.get("nombre", "")),
        ("Tipo de documento", d.get("tipo_doc", "")),
        ("Número de identificación", d.get("num_doc", "")),
        ("Contacto telefónico", d.get("telefono", "")),
    ]
    for i, (lab, val) in enumerate(labels_r0):
        write_cell(t1.cell(0, i), lab, size=7, bold=True, align="center", shade=LIGHT_GREEN)
        write_cell(t1.cell(1, i), val or "________________", size=8, align="center", shade=YELLOW if not val else WHITE)

    labels_r1 = [
        ("Correo electrónico institucional", d.get("correo_inst", "")),
        ("Correo electrónico personal", d.get("correo_pers", "")),
        ("Dirección de residencia", d.get("direccion", "")),
        ("Número de grupo", d.get("grupo", "")),
    ]
    for i, (lab, val) in enumerate(labels_r1):
        write_cell(t1.cell(2, i), lab, size=7, bold=True, align="center", shade=LIGHT_GREEN)
        write_cell(t1.cell(3, i), val or "________________", size=8, align="center", shade=YELLOW if not val else WHITE)

    t2 = doc.add_table(rows=2, cols=3)
    fields = [
        ("Modalidad de formación", d.get("modalidad_formacion", "")),
        ("Programa de formación", d.get("programa", "")),
        ("Modalidad de ejecución de la etapa productiva\n(presencial o virtual)", d.get("modalidad_etapa", "")),
    ]
    for i, (lab, val) in enumerate(fields):
        write_cell(t2.cell(0, i), lab, size=7, bold=True, align="center", shade=LIGHT_GREEN)
        write_cell(t2.cell(1, i), val or "________________", size=8, align="center", shade=YELLOW if not val else WHITE)

    t3 = doc.add_table(rows=2, cols=2)
    write_cell(
        t3.cell(0, 0),
        "¿Realiza la etapa productiva con una entidad u organización en el exterior? (sí o no)",
        size=7,
        bold=True,
        align="center",
        shade=LIGHT_GREEN,
    )
    write_cell(t3.cell(0, 1), "País donde realiza la etapa productiva", size=7, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t3.cell(1, 0), d.get("exterior", "No") or "No", size=8, align="center", shade=YELLOW)
    write_cell(t3.cell(1, 1), d.get("pais", "Colombia") or "Colombia", size=8, align="center", shade=YELLOW)

    doc.add_paragraph()
    add_heading_para(doc, "Datos de la entidad co-formadora", size=10, align="left", space_after=4)
    t4 = doc.add_table(rows=2, cols=3)
    emp = [
        ("Nombre de la entidad, empresa, institución u organización", d.get("empresa", "")),
        ("NIT", d.get("nit", "")),
        ("Dirección de la entidad", d.get("dir_empresa", "")),
    ]
    for i, (lab, val) in enumerate(emp):
        write_cell(t4.cell(0, i), lab, size=7, bold=True, align="center", shade=LIGHT_GREEN)
        write_cell(t4.cell(1, i), val or "________________", size=8, align="center", shade=YELLOW if not val else WHITE)
    doc.add_paragraph()


def build_coformador(doc, d):
    add_heading_para(
        doc,
        "Datos de la persona encargada del proceso formativo del aprendiz en la entidad co-formadora",
        size=10,
        align="left",
        space_after=4,
    )
    t = doc.add_table(rows=2, cols=4)
    fields = [
        ("Nombre completo del ente co-formador\n(Jefe inmediato / Supervisor)", d.get("jefe_nombre", "")),
        ("Cargo del ente co-formador", d.get("jefe_cargo", "")),
        ("Contacto telefónico del ente co-formador", d.get("jefe_tel", "")),
        ("Correo electrónico del ente co-formador", d.get("jefe_correo", "")),
    ]
    for i, (lab, val) in enumerate(fields):
        write_cell(t.cell(0, i), lab, size=7, bold=True, align="center", shade=LIGHT_GREEN)
        write_cell(t.cell(1, i), val or "________________", size=8, align="center", shade=YELLOW if not val else WHITE)
    doc.add_paragraph()


def build_instructor(doc, d):
    add_heading_para(doc, "Datos del instructor de seguimiento", size=10, align="left", space_after=4)
    t = doc.add_table(rows=2, cols=2)
    write_cell(t.cell(0, 0), "Nombre completo del instructor de seguimiento", size=7, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(0, 1), "Correo electrónico del instructor de seguimiento", size=7, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(1, 0), d.get("instructor_nombre", "") or "________________", size=8, align="center", shade=YELLOW)
    write_cell(t.cell(1, 1), d.get("instructor_correo", "") or "________________", size=8, align="center", shade=YELLOW)
    doc.add_paragraph()


def build_alternativa(doc, d):
    add_heading_para(
        doc,
        'Seleccione con una "X" la alternativa de etapa productiva que está realizando',
        size=10,
        align="left",
        space_after=4,
    )
    alternativas = [
        "Contrato de aprendizaje",
        "Monitoria",
        "Proyecto productivo",
        "Contrato de vínculo formativo",
        "Vínculo laboral",
    ]
    marcada = d.get("alternativa", "")
    t = doc.add_table(rows=len(alternativas) + 1, cols=2)
    write_cell(t.cell(0, 0), "Alternativa de etapa productiva", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(0, 1), 'Marque con una "X"', size=8, bold=True, align="center", shade=LIGHT_GREEN)
    for i, alt in enumerate(alternativas, start=1):
        write_cell(t.cell(i, 0), alt, size=8, align="left")
        marca = "X" if marcada.lower() == alt.lower() else ""
        write_cell(t.cell(i, 1), marca or "☐", size=10, bold=True, align="center", shade=YELLOW if marca else WHITE)
    doc.add_paragraph()


def build_actividades(doc, actividades):
    add_heading_para(doc, "Descripción de las actividades realizadas", size=10, align="left", space_after=4)
    headers = [
        "Descripción de la actividad",
        "Competencias del programa de formación aplicadas",
        "Fecha de inicio\n(dd/mm/aa)",
        "Fecha de fin\n(dd/mm/aa)",
        "Evidencia de cumplimiento\n(documento, proceso, producto, entregable u otro)",
        "Observaciones, inasistencias, dificultades y/o comentarios",
    ]
    rows = max(len(actividades), 5)
    t = doc.add_table(rows=rows + 1, cols=6)
    for i, h in enumerate(headers):
        write_cell(t.cell(0, i), h, size=7, bold=True, align="center", shade=LIGHT_GREEN)

    for r in range(rows):
        act = actividades[r] if r < len(actividades) else {}
        vals = [
            act.get("descripcion", ""),
            act.get("competencias", ""),
            act.get("fecha_inicio", ""),
            act.get("fecha_fin", ""),
            act.get("evidencia", ""),
            act.get("observaciones", ""),
        ]
        for c, v in enumerate(vals):
            write_cell(
                t.cell(r + 1, c),
                v or (" " * 8),
                size=7,
                align="left" if c in (0, 1, 4, 5) else "center",
                shade=YELLOW if not v else WHITE,
            )
        # altura mínima de fila
        t.rows[r + 1].height = Cm(1.6)
        t.rows[r + 1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    doc.add_paragraph()


def build_arl(doc, d):
    add_heading_para(doc, "Información afiliación a la ARL", size=10, align="left", space_after=4)
    p = doc.add_paragraph()
    run = p.add_run(
        "Decreto 055 de 2015, por el cual se reglamenta la afiliación de estudiantes al Sistema General "
        "de Riesgos Laborales y se dictan otras disposiciones. Este espacio debe ser siempre diligenciado."
    )
    set_run_font(run, size=7, bold=False)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Artículo 11. Obligaciones de la institución de educación:\n"
        "1. Revisar periódicamente que el estudiante en práctica desarrolle labores relacionadas "
        "exclusivamente con su programa de formación.\n"
        "2. Verificar que el espacio de práctica cuente con los elementos de protección personal "
        "apropiados según el riesgo ocupacional."
    )
    set_run_font(run2, size=7)
    p2.paragraph_format.space_after = Pt(6)

    t = doc.add_table(rows=2, cols=4)
    headers = [
        "¿El aprendiz se encuentra afiliado a la ARL?",
        "Indique el nivel de riesgo actual",
        "¿El nivel de riesgo de la ARL corresponde a las actividades que desarrolla el aprendiz en la empresa? (SI / NO)",
        "¿El aprendiz cuenta con los EPP requeridos para desarrollar su etapa productiva? (SI / NO / NA)",
    ]
    vals = [
        d.get("arl_afiliado", ""),
        d.get("arl_nivel", ""),
        d.get("arl_corresponde", ""),
        d.get("arl_epp", ""),
    ]
    for i, h in enumerate(headers):
        write_cell(t.cell(0, i), h, size=7, bold=True, align="center", shade=LIGHT_GREEN)
        write_cell(t.cell(1, i), vals[i] or "________________", size=8, align="center", shade=YELLOW if not vals[i] else WHITE)
    doc.add_paragraph()


def build_firmas(doc, d):
    p = doc.add_paragraph()
    run = p.add_run(
        "Aprendiz: recuerde diligenciar completamente el formato de bitácora y entregarlo o cargarlo "
        "al espacio asignado para este."
    )
    set_run_font(run, size=8, bold=True)
    p.paragraph_format.space_after = Pt(10)

    t = doc.add_table(rows=4, cols=3)
    write_cell(t.cell(0, 0), "Firma del aprendiz", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(0, 1), "Fecha entrega bitácora", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(0, 2), "", size=8, shade=LIGHT_GREEN)
    write_cell(t.cell(1, 0), "\n\n________________________\n", size=8, align="center")
    write_cell(t.cell(1, 1), d.get("fecha_entrega", "") or "____/____/________", size=9, align="center", shade=YELLOW)
    write_cell(t.cell(1, 2), "", size=8)

    write_cell(t.cell(2, 0), "Firma del instructor de seguimiento", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(2, 1), "Firma del ente co-formador", size=8, bold=True, align="center", shade=LIGHT_GREEN)
    write_cell(t.cell(2, 2), "", size=8, shade=LIGHT_GREEN)
    write_cell(t.cell(3, 0), "\n\n________________________\n", size=8, align="center")
    write_cell(t.cell(3, 1), "\n\n________________________\n", size=8, align="center")
    write_cell(t.cell(3, 2), "", size=8)

    doc.add_paragraph()
    nota = doc.add_paragraph()
    run = nota.add_run(
        "Nota: Con el diligenciamiento de este formato autorizo al SENA para la recolección y tratamiento "
        "de mis datos personales, conforme a la política de datos personales de la entidad GOR-POL-006. "
        "Entiendo que los datos serán objeto de recolección, almacenamiento, uso, circulación, supresión, "
        "transferencia, transmisión, cesión y todo el tratamiento, realizados por el SENA."
    )
    set_run_font(run, size=7, bold=False)
    nota.paragraph_format.space_before = Pt(8)

    anexo = doc.add_paragraph()
    run_a = anexo.add_run(
        "Anexo (opcional): Es opcional relacionar evidencia fotográfica de las actividades desarrolladas. "
        "(No aplica documentos de la empresa u otros aspectos sensibles)."
    )
    set_run_font(run_a, size=8, bold=True)
    anexo.paragraph_format.space_before = Pt(8)


def set_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)


def generar(datos, actividades, salida):
    doc = Document()
    set_page(doc)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    build_header(doc)
    build_bitacora_periodo(doc, datos)
    build_datos_aprendiz(doc, datos)
    build_coformador(doc, datos)
    build_instructor(doc, datos)
    build_alternativa(doc, datos)
    build_actividades(doc, actividades)
    build_arl(doc, datos)
    build_firmas(doc, datos)

    doc.save(salida)
    print(f"Generado: {salida}")


if __name__ == "__main__":
    from datos_bitacora import DATOS, ACTIVIDADES

    n = DATOS.get("bitacora_n", "X")
    generar(
        DATOS,
        ACTIVIDADES,
        f"/workspace/Bitacora{n}_Esteban_Andres_Perdomo_Rojas.docx",
    )
